"""Excel parsing, Word→PDF rendering with placeholders, and PDF encryption."""
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import List, Dict, Tuple, Optional

from openpyxl import load_workbook
from docxtpl import DocxTemplate
import pikepdf


def parse_excel(file_path: str) -> Tuple[List[str], List[Dict[str, str]]]:
    """Parse Excel: first row=headers (col1=email, col2=password), rest=data.
    Returns (headers, rows) where rows is list of dicts keyed by header name.
    """
    wb = load_workbook(file_path, data_only=True)
    ws = wb.active
    rows_iter = ws.iter_rows(values_only=True)
    try:
        header_row = next(rows_iter)
    except StopIteration:
        return [], []
    headers = [str(h).strip() if h is not None else f"col_{i}" for i, h in enumerate(header_row)]
    data: List[Dict[str, str]] = []
    for row in rows_iter:
        if row is None or all(v is None or str(v).strip() == "" for v in row):
            continue
        record = {}
        for i, h in enumerate(headers):
            v = row[i] if i < len(row) else None
            record[h] = "" if v is None else str(v)
        data.append(record)
    return headers, data


def replace_placeholders_text(text: str, data: Dict[str, str]) -> str:
    """Replace {field} placeholders in plain text/HTML strings."""
    def _repl(m):
        key = m.group(1).strip()
        return str(data.get(key, m.group(0)))
    return re.sub(r"\{([^{}]+)\}", _repl, text or "")


def render_docx_to_pdf(template_path: str, context: Dict[str, str], output_dir: str) -> str:
    """Render docxtpl template with context, convert to PDF via LibreOffice. Returns PDF path."""
    os.makedirs(output_dir, exist_ok=True)
    tmp_docx = os.path.join(output_dir, "rendered.docx")

    # docxtpl uses Jinja2 syntax. Our placeholders are {field} (single braces).
    # We need to convert {field} to {{ field }} before rendering with docxtpl.
    # Strategy: load with python-docx, replace text in runs and tables, then save.
    from docx import Document
    doc = Document(template_path)

    def _sub(text: str) -> str:
        return replace_placeholders_text(text, context)

    # Replace in paragraphs (preserving runs as best as possible)
    for para in doc.paragraphs:
        _replace_in_paragraph(para, context)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, context)
    # Headers / footers
    for section in doc.sections:
        for hdr in [section.header, section.footer]:
            for para in hdr.paragraphs:
                _replace_in_paragraph(para, context)

    doc.save(tmp_docx)

    # Convert to PDF via LibreOffice headless
    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, tmp_docx],
        capture_output=True, text=True, timeout=120,
    )
    pdf_path = os.path.join(output_dir, "rendered.pdf")
    if not os.path.exists(pdf_path):
        raise RuntimeError(f"PDF conversion failed: {result.stderr or result.stdout}")
    return pdf_path


def _replace_in_paragraph(paragraph, context: Dict[str, str]):
    """Replace placeholders in a paragraph while preserving formatting where possible."""
    full_text = paragraph.text
    if "{" not in full_text:
        return
    new_text = replace_placeholders_text(full_text, context)
    if new_text == full_text:
        return
    # Simple approach: keep first run's formatting, replace its text, clear others
    if not paragraph.runs:
        return
    first = paragraph.runs[0]
    first.text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def encrypt_pdf(input_pdf: str, output_pdf: str, password: str) -> None:
    """Encrypt PDF with AES-256 using pikepdf."""
    with pikepdf.open(input_pdf) as pdf:
        pdf.save(
            output_pdf,
            encryption=pikepdf.Encryption(owner=password, user=password, R=6),
        )


def build_personalized_pdf(template_path: str, row_data: Dict[str, str], password: Optional[str], work_dir: str, out_name: str = "document.pdf") -> str:
    """Build a per-recipient PDF, encrypt if password provided. Returns final PDF path."""
    rendered_dir = tempfile.mkdtemp(dir=work_dir)
    try:
        rendered_pdf = render_docx_to_pdf(template_path, row_data, rendered_dir)
        final_path = os.path.join(work_dir, out_name)
        if password and password.strip():
            encrypt_pdf(rendered_pdf, final_path, password.strip())
        else:
            shutil.copy(rendered_pdf, final_path)
        return final_path
    finally:
        shutil.rmtree(rendered_dir, ignore_errors=True)
